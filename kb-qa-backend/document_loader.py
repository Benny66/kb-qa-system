"""
统一文档解析模块。
当前支持：
- txt
- md / markdown
- pdf
- docx
"""

from __future__ import annotations

import os

from docx import Document
from pypdf import PdfReader

TEXT_EXTENSIONS = {"txt", "md", "markdown"}
PDF_EXTENSIONS = {"pdf"}
WORD_EXTENSIONS = {"docx"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | PDF_EXTENSIONS | WORD_EXTENSIONS


def get_file_extension(filename: str) -> str:
    return filename.rsplit(".", 1)[1].lower() if "." in filename else ""


def is_supported_document(filename: str) -> bool:
    return get_file_extension(filename) in SUPPORTED_EXTENSIONS


def read_text_document(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="gbk", errors="replace") as f:
            return f.read()


def read_pdf_document(file_path: str) -> str:
    reader = PdfReader(file_path)
    parts: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts).strip()


def read_docx_document(file_path: str) -> str:
    document = Document(file_path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def extract_document_text(file_path: str, filename: str | None = None) -> str:
    """根据扩展名自动解析文档内容。"""
    target_name = filename or os.path.basename(file_path)
    ext = get_file_extension(target_name)

    if ext in TEXT_EXTENSIONS:
        return read_text_document(file_path)
    if ext in PDF_EXTENSIONS:
        return read_pdf_document(file_path)
    if ext in WORD_EXTENSIONS:
        return read_docx_document(file_path)

    raise ValueError(f"暂不支持的文件格式：.{ext}" if ext else "无法识别文件扩展名")
